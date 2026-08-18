"""
TalentConnect AI Agent — utilities.
File text extraction, OpenAI wrapper, and a simple in-session record store
with JSON export/import (Streamlit Community Cloud storage is ephemeral,
so export is the durable save).
"""

import io
import json
import os
from datetime import datetime

import streamlit as st


# ---------------------------------------------------------------------------
# OpenAI
# ---------------------------------------------------------------------------

def get_api_key() -> str | None:
    """Read the API key from Streamlit secrets or environment."""
    try:
        key = st.secrets.get("OPENAI_API_KEY", None)
    except Exception:
        key = None
    return key or os.environ.get("OPENAI_API_KEY")


def get_model() -> str:
    try:
        model = st.secrets.get("OPENAI_MODEL", None)
    except Exception:
        model = None
    return model or os.environ.get("OPENAI_MODEL") or "gpt-4o-mini"


def _configured_password() -> str | None:
    """The app password, read from secrets or environment. If none is set,
    the gate is disabled (app is open) — so nothing breaks before setup."""
    try:
        pw = st.secrets.get("APP_PASSWORD", None)
    except Exception:
        pw = None
    return pw or os.environ.get("APP_PASSWORD")


def check_password() -> bool:
    """Simple shared-password gate. Returns True if access is granted.

    Behaviour:
      - If no APP_PASSWORD is configured, the gate is OFF (returns True) so the
        app still runs during setup / local testing.
      - Otherwise the user must enter the password once per session.
    This protects the candidate database from being read by anyone with just the
    public link. It is a shared team password, not per-user accounts.
    """
    password = _configured_password()
    if not password:
        return True  # gate not configured yet — app open

    if st.session_state.get("_authed"):
        return True

    st.markdown("## 🔒 TalentConnect AI Agent")
    st.caption("This tool contains candidate information. Please enter the team "
               "password to continue.")
    entered = st.text_input("Password", type="password", key="_pw_input")
    if st.button("Enter", type="primary"):
        if entered == password:
            st.session_state["_authed"] = True
            st.rerun()
        else:
            st.error("Incorrect password. Please try again.")
    st.caption("Access is limited to the Lithan Talent Management team.")
    return False


def run_ai(system_prompt: str, user_content: str, temperature: float = 0.4) -> str:
    """Call the OpenAI chat completions API and return the text response.

    Raises RuntimeError with a human-readable message on failure so the UI
    can show it directly.
    """
    api_key = get_api_key()
    if not api_key:
        raise RuntimeError(
            "No OpenAI API key found. Add OPENAI_API_KEY in the app's Secrets "
            "(Streamlit Cloud: Manage app → Settings → Secrets)."
        )

    try:
        from openai import OpenAI

        client = OpenAI(api_key=api_key)
        response = client.chat.completions.create(
            model=get_model(),
            temperature=temperature,
            max_tokens=8000,  # allow long outputs (e.g. a 50-candidate match report)
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_content},
            ],
        )
        return response.choices[0].message.content or ""
    except Exception as exc:  # surface a readable error in the UI
        raise RuntimeError(f"OpenAI request failed: {exc}") from exc


def research_company(company: str, job_title: str = "") -> str:
    """Use the OpenAI web-search tool to research a company so matching can
    judge fit against the real employer, not just the JD text. Returns a short
    factual brief, or an empty string if research is unavailable/fails (the
    caller then proceeds with JD-only matching)."""
    api_key = get_api_key()
    if not api_key or not company.strip():
        return ""
    try:
        from openai import OpenAI

        client = OpenAI(api_key=api_key)
        prompt = (
            f"Research the company \"{company}\" (Singapore context if applicable). "
            f"The role being hired is: {job_title or 'not specified'}.\n\n"
            "Give a short factual brief (under 200 words) covering: what the company "
            "does (industry, products/services), its size/stage if known, the sectors "
            "or clients it serves, and any context that affects what kind of hire would "
            "fit (e.g. fast-paced startup vs established MNC, technical vs commercial "
            "focus). If you cannot find reliable information, say so plainly and do not "
            "invent details."
        )
        # Try the web-search-enabled Responses API first
        try:
            resp = client.responses.create(
                model=get_model(),
                tools=[{"type": "web_search_preview"}],
                input=prompt,
            )
            text = getattr(resp, "output_text", "") or ""
            if text.strip():
                return text.strip()
        except Exception:
            pass  # fall through to non-search fallback
        # Fallback: model's own knowledge, clearly caveated
        resp2 = client.chat.completions.create(
            model=get_model(), temperature=0.3, max_tokens=400,
            messages=[
                {"role": "system", "content": "You provide brief, factual company "
                 "briefs. If unsure, say so; never invent specifics."},
                {"role": "user", "content": prompt},
            ],
        )
        out = resp2.choices[0].message.content or ""
        return (out.strip() + "\n\n(Note: based on general knowledge, not a live web "
                "lookup.)") if out.strip() else ""
    except Exception:
        return ""  # research is best-effort; matching continues without it


# ---------------------------------------------------------------------------
# File reading (PDF / DOCX / TXT)
# ---------------------------------------------------------------------------

def extract_text(uploaded_file) -> str:
    """Extract plain text from a Streamlit UploadedFile (pdf, docx, txt)."""
    if uploaded_file is None:
        return ""

    name = uploaded_file.name.lower()
    data = uploaded_file.read()
    uploaded_file.seek(0)

    if name.endswith(".pdf"):
        return _extract_pdf(data)
    if name.endswith(".docx"):
        return _extract_docx(data)
    # default: treat as text
    for encoding in ("utf-8", "utf-16", "gb18030", "latin-1"):
        try:
            return data.decode(encoding)
        except (UnicodeDecodeError, LookupError):
            continue
    return data.decode("utf-8", errors="replace")


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    pages = [(page.extract_text() or "") for page in reader.pages]
    text = "\n\n".join(pages).strip()
    if not text:
        return "[This PDF appears to be scanned images with no extractable text. Please paste the text manually.]"
    return text


def _extract_docx(data: bytes) -> str:
    import docx

    document = docx.Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(part for part in parts if part.strip())


def read_input(label: str, key: str, height: int = 220) -> str:
    """Render an upload-or-paste input block and return the resulting text."""
    uploaded = st.file_uploader(
        f"Upload {label} (PDF, DOCX or TXT)",
        type=["pdf", "docx", "txt"],
        key=f"{key}_file",
    )
    pasted = st.text_area(
        f"…or paste {label} text here",
        key=f"{key}_text",
        height=height,
        placeholder="Paste text here if you are not uploading a file.",
    )
    if uploaded is not None:
        try:
            text = extract_text(uploaded)
            st.caption(f"Read {len(text):,} characters from {uploaded.name}")
            return text
        except Exception as exc:
            st.error(f"Could not read {uploaded.name}: {exc}")
            return ""
    return pasted or ""


# ---------------------------------------------------------------------------
# Data store: candidates, jobs, generated outputs (session + JSON backup)
# ---------------------------------------------------------------------------

OUTPUT_TYPES = [
    "Candidate profile", "Coaching notes", "WhatsApp follow-up", "JD analysis",
    "Match report", "Outreach", "Fitment analysis", "Profile submission", "Interview prep",
]


def _now() -> str:
    return datetime.now().strftime("%Y-%m-%d %H:%M")


def _new_id(prefix: str) -> str:
    return f"{prefix}-{datetime.now().strftime('%Y%m%d%H%M%S%f')}"


STORE_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                          "talentconnect_data.json")


def init_store() -> None:
    if st.session_state.get("_store_loaded"):
        return
    st.session_state.setdefault("candidates", [])
    st.session_state.setdefault("jobs", [])
    st.session_state.setdefault("outputs", [])
    st.session_state.setdefault("status_overrides", {})
    # Auto-load persisted data so records survive page refreshes
    if os.path.exists(STORE_FILE):
        try:
            with open(STORE_FILE, encoding="utf-8") as fh:
                data = json.load(fh)
            st.session_state.candidates = data.get("candidates", [])
            st.session_state.jobs = data.get("jobs", [])
            st.session_state.outputs = data.get("outputs", [])
            st.session_state.status_overrides = data.get("status_overrides", {})
        except Exception:
            pass  # corrupted/missing file: start clean rather than crash
    st.session_state["_store_loaded"] = True


def persist() -> None:
    """Write the store to disk after every app run, so any change made
    during the run is saved automatically and survives page refreshes."""
    if not st.session_state.get("_store_loaded"):
        return
    try:
        with open(STORE_FILE, "w", encoding="utf-8") as fh:
            fh.write(export_store())
    except Exception:
        pass


def add_candidate(name: str, resume: str = "", salary: str = "", notice: str = "") -> dict:
    init_store()
    candidate = {
        "id": _new_id("cand"), "name": name.strip(), "resume": resume,
        "salary": salary, "notice": notice, "profile": "", "coaching_notes": "",
        "created": _now(),
    }
    st.session_state.candidates.append(candidate)
    return candidate


def get_candidate(cid: str) -> dict | None:
    init_store()
    return next((c for c in st.session_state.candidates if c["id"] == cid), None)


def db_candidate_as_resume(rec: dict) -> str:
    """Render a bundled database record as résumé-style text the AI modules
    (matching, profile, fitment) can consume, so a pre-loaded candidate works
    exactly like a manually added one."""
    lines = [
        f"Name: {rec.get('full_name','')}",
        f"Talent Specialist: {rec.get('specialist','')}  |  "
        f"Course: {rec.get('course_code','')} ({rec.get('course_name','')})  |  "
        f"Cohort: {rec.get('cohort','')}",
        f"Experience: {rec.get('years_experience','')}",
        f"Domain knowledge: {rec.get('domain','')}",
        f"Industry background: {rec.get('industry_background','')}",
        "",
        f"Summary: {rec.get('prior_experience_summary','')}",
        "",
        f"Skills: {', '.join(rec.get('skills', []))}",
        "",
        f"Skill Marriage (unique value): {rec.get('skill_marriage','')}",
        f"Recommended roles: {', '.join(rec.get('recommended_roles', []))}",
        f"Seniority: {rec.get('seniority','')} — {rec.get('seniority_note','')}",
    ]
    return "\n".join(lines)


def candidate_status(rec: dict) -> str:
    """Return a database candidate's current status. In-app changes are stored
    as overrides in the persistent store (keyed by full name); if none, fall
    back to the baseline 'status' in candidate_db.py, else 'Active'."""
    init_store()
    overrides = st.session_state.setdefault("status_overrides", {})
    name = rec.get("full_name", "").strip().lower()
    if name in overrides:
        return overrides[name]
    return rec.get("status", "Active")


def set_candidate_status(full_name: str, status: str) -> None:
    """Set a database candidate's status (Active / Placed / Inactive) and
    persist it. Works even though candidate_db.py itself is read-only."""
    init_store()
    overrides = st.session_state.setdefault("status_overrides", {})
    overrides[full_name.strip().lower()] = status
    persist()


def import_db_candidate(rec: dict) -> dict:
    """Copy a bundled database record into the user's working candidate list
    so it flows through coaching, matching, fitment and outreach. Returns the
    existing record if this person was already imported (matched by name)."""
    init_store()
    existing = next((c for c in st.session_state.candidates
                     if c["name"].strip().lower() == rec["full_name"].strip().lower()), None)
    if existing:
        return existing
    candidate = {
        "id": _new_id("cand"), "name": rec["full_name"].strip(),
        "resume": db_candidate_as_resume(rec),
        "salary": "", "notice": "",
        "profile": "", "coaching_notes": "",
        "email": rec.get("email", ""),
        "source": "database", "specialist": rec.get("specialist", ""),
        "course_code": rec.get("course_code", ""), "cohort": rec.get("cohort", ""),
        "created": _now(),
    }
    st.session_state.candidates.append(candidate)
    return candidate


def add_job(title: str, employer: str = "", jd: str = "") -> dict:
    init_store()
    job = {
        "id": _new_id("job"), "title": title.strip(), "employer": employer.strip(),
        "jd": jd, "analysis": "", "created": _now(),
    }
    st.session_state.jobs.append(job)
    return job


def get_job(jid: str) -> dict | None:
    init_store()
    return next((j for j in st.session_state.jobs if j["id"] == jid), None)


def set_job_shortlist(job_id: str, candidate_ids: list) -> None:
    """Associate a set of candidates with a job as its shortlist. Merges with
    any existing shortlist (no duplicates) and persists."""
    init_store()
    job = get_job(job_id)
    if not job:
        return
    current = job.get("shortlist", [])
    for cid in candidate_ids:
        if cid not in current:
            current.append(cid)
    job["shortlist"] = current
    persist()


def job_shortlist_names(job: dict) -> list:
    """Return the display names of a job's shortlisted candidates."""
    init_store()
    by_id = {c["id"]: c["name"] for c in st.session_state.candidates}
    return [by_id.get(cid, "(removed)") for cid in job.get("shortlist", [])]


def add_output(output_type: str, content: str, candidate_id: str | None = None,
               job_id: str | None = None, title: str = "") -> dict:
    """Auto-record every generated artefact, linked to candidate and/or job."""
    init_store()
    record = {
        "id": _new_id("out"), "type": output_type, "title": title or output_type,
        "content": content, "candidate_id": candidate_id, "job_id": job_id,
        "created": _now(),
    }
    st.session_state.outputs.append(record)
    return record


def outputs_for(candidate_id: str | None = None, job_id: str | None = None,
                output_type: str | None = None) -> list[dict]:
    init_store()
    results = st.session_state.outputs
    if candidate_id:
        results = [o for o in results if o.get("candidate_id") == candidate_id]
    if job_id:
        results = [o for o in results if o.get("job_id") == job_id]
    if output_type:
        results = [o for o in results if o.get("type") == output_type]
    return results


def delete_output(output_id: str) -> None:
    init_store()
    st.session_state.outputs = [o for o in st.session_state.outputs if o["id"] != output_id]


def delete_candidate(cid: str) -> None:
    init_store()
    st.session_state.candidates = [c for c in st.session_state.candidates if c["id"] != cid]
    st.session_state.outputs = [o for o in st.session_state.outputs if o.get("candidate_id") != cid]


def delete_job(jid: str) -> None:
    init_store()
    st.session_state.jobs = [j for j in st.session_state.jobs if j["id"] != jid]
    st.session_state.outputs = [o for o in st.session_state.outputs if o.get("job_id") != jid]


def export_store() -> str:
    init_store()
    return json.dumps(
        {
            "candidates": st.session_state.candidates,
            "jobs": st.session_state.jobs,
            "outputs": st.session_state.outputs,
            "status_overrides": st.session_state.get("status_overrides", {}),
        },
        ensure_ascii=False, indent=2,
    )


def import_store(json_text: str) -> dict:
    init_store()
    data = json.loads(json_text)
    counts = {}
    for key in ("candidates", "jobs", "outputs"):
        items = data.get(key, [])
        existing = {x["id"] for x in st.session_state[key]}
        added = [x for x in items if isinstance(x, dict) and x.get("id") and x["id"] not in existing]
        st.session_state[key].extend(added)
        counts[key] = len(added)
    # merge status overrides (import wins for any placed/inactive marks)
    if isinstance(data.get("status_overrides"), dict):
        st.session_state.setdefault("status_overrides", {}).update(data["status_overrides"])
    return counts


# ---------------------------------------------------------------------------
# Word (.docx) export of generated markdown documents
# ---------------------------------------------------------------------------

def markdown_to_docx_bytes(md_text: str) -> bytes:
    """Convert generated markdown to a .docx file. Prefers pandoc (installed
    on Streamlit Cloud via packages.txt); falls back to a python-docx
    converter that handles headings, bullets, bold and tables."""
    try:
        return _pandoc_docx(md_text)
    except Exception:
        return _pydocx_docx(md_text)


def _pandoc_docx(md_text: str) -> bytes:
    import subprocess
    import tempfile

    with tempfile.TemporaryDirectory() as tmp:
        src = os.path.join(tmp, "doc.md")
        out = os.path.join(tmp, "doc.docx")
        with open(src, "w", encoding="utf-8") as fh:
            fh.write(md_text)
        subprocess.run(["pandoc", src, "-f", "markdown", "-t", "docx", "-o", out],
                       check=True, timeout=60, capture_output=True)
        with open(out, "rb") as fh:
            return fh.read()


def _strip_md(text: str) -> str:
    text = text.replace("**", "").replace("__", "")
    if text.startswith("*") and text.endswith("*") and len(text) > 2:
        text = text[1:-1]
    return text.strip()


def _pydocx_docx(md_text: str) -> bytes:
    import docx
    from docx.shared import Pt

    document = docx.Document()
    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()
        if not stripped:
            i += 1
            continue
        # Tables: a header row followed by a |---| separator
        if (stripped.startswith("|") and i + 1 < len(lines)
                and set(lines[i + 1].strip()) <= set("|-: ")
                and "-" in lines[i + 1]):
            rows = []
            header = [_strip_md(c) for c in stripped.strip("|").split("|")]
            rows.append(header)
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append([_strip_md(c) for c in lines[i].strip().strip("|").split("|")])
                i += 1
            ncols = max(len(r) for r in rows)
            table = document.add_table(rows=len(rows), cols=ncols)
            table.style = "Table Grid"
            for r, row in enumerate(rows):
                for c in range(ncols):
                    cell_text = row[c] if c < len(row) else ""
                    cell = table.cell(r, c)
                    cell.text = cell_text
                    if r == 0:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
            continue
        if stripped.startswith("#"):
            level = min(len(stripped) - len(stripped.lstrip("#")), 4)
            document.add_heading(_strip_md(stripped.lstrip("#").strip()), level=level)
        elif stripped.startswith(("- ", "* ", "• ")):
            document.add_paragraph(_strip_md(stripped[2:]), style="List Bullet")
        elif stripped[:2].isdigit() or (stripped[0].isdigit() and stripped[1] in ".)"):
            document.add_paragraph(_strip_md(stripped.split(" ", 1)[-1]), style="List Number")
        elif stripped == "---":
            pass
        else:
            paragraph = document.add_paragraph()
            # Bold whole-line **Label:** value patterns simply by stripping markers
            run = paragraph.add_run(_strip_md(stripped))
            run.font.size = Pt(11)
        i += 1

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
